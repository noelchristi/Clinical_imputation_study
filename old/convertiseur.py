# -*- coding: utf-8 -*-
"""
Convertisseur Markdown + LaTeX math -> DOCX

Stratégie :
  1. pandoc (natif OMML) — meilleur rendu mathématique dans Word.
  2. pypandoc (fallback Python).
  3. Rendu matplotlib : transforme les équations LaTeX en images PNG et
     les insère dans un .docx via python-docx (fallback ultime).

Usage :
  python convertiseur.py article_scientifique.md article_scientifique.docx
"""

import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

# ---------------------------------------------------------------------------
# Détection pandoc
# ---------------------------------------------------------------------------

def check_pandoc_installed():
    try:
        subprocess.run(
            ["pandoc", "--version"],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True
        )
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False


# ---------------------------------------------------------------------------
# Expressions régulières pour le parsing mathématique
# ---------------------------------------------------------------------------

MATH_BLOCK_RE = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
MATH_INLINE_RE = re.compile(r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)")


# ---------------------------------------------------------------------------
# 1. Conversion pandoc
# ---------------------------------------------------------------------------

def convert_with_pandoc(md_path: str, docx_path: str):
    cmd = [
        "pandoc",
        md_path,
        "-o", docx_path,
        "--from=markdown+tex_math_dollars+pipe_tables+fenced_code_blocks",
        "--to=docx",
        "--resource-path", str(Path(md_path).parent),
    ]
    subprocess.run(cmd, check=True)
    return True


# ---------------------------------------------------------------------------
# 2. Conversion via pypandoc
# ---------------------------------------------------------------------------

def convert_with_pypandoc(md_path: str, docx_path: str):
    import pypandoc
    pypandoc.convert_file(
        md_path, 'docx',
        outputfile=docx_path,
        extra_args=['--from=markdown+tex_math_dollars+pipe_tables+fenced_code_blocks'],
    )
    return True


# ---------------------------------------------------------------------------
# 3. Fallback : rendu matplotlib + python-docx
# ---------------------------------------------------------------------------

def _render_latex_png(latex: str, out_path: str, dpi: int = 200):
    """Rendu d'une équation LaTeX en PNG via matplotlib."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    # Essayer mathtext (léger, pas besoin de LaTeX système)
    try:
        fig, ax = plt.subplots(figsize=(0.01, 0.01))
        text = ax.text(0.5, 0.5, f"${latex}$", fontsize=12,
                       ha='center', va='center',
                       transform=ax.transAxes)
        ax.axis('off')
        fig.savefig(out_path, dpi=dpi, bbox_inches='tight',
                    pad_inches=0.05, transparent=False, facecolor='white')
        plt.close(fig)
        return True
    except Exception:
        plt.close('all')
        raise


def _extract_placeholders(text):
    """Remplace les blocs math par des placeholders et retourne text + dicts."""
    math_blocks = {}
    math_inlines = {}

    def block_repl(m):
        key = f"\x00BLOCK{len(math_blocks)}\x00"
        math_blocks[key] = m.group(1).strip()
        return key

    def inline_repl(m):
        key = f"\x00INLINE{len(math_inlines)}\x00"
        math_inlines[key] = m.group(1).strip()
        return key

    text = MATH_BLOCK_RE.sub(block_repl, text)
    text = MATH_INLINE_RE.sub(inline_repl, text)
    return text, math_blocks, math_inlines


def _add_formatted_paragraph(doc, text, math_blocks, math_inlines, tmpdir):
    """Ajoute un paragraphe au document en gérant bold, italic, code, math."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()

    # Split en tokens: **bold**, *italic*, `code`, placeholders
    token_pat = re.compile(
        r'(\*\*(.+?)\*\*|'
        r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|'
        r'`(.+?)`|'
        r'(\x00(?:BLOCK|INLINE)\d+\x00))'
    )

    pos = 0
    for m in token_pat.finditer(text):
        # Texte avant le token
        before = text[pos:m.start()]
        if before:
            p.add_run(before)

        full = m.group(0)
        bold_content = m.group(2)
        italic_content = m.group(3)
        code_content = m.group(4)
        placeholder = m.group(5)

        if bold_content is not None:
            run = p.add_run(bold_content)
            run.bold = True
        elif italic_content is not None:
            run = p.add_run(italic_content)
            run.italic = True
        elif code_content is not None:
            run = p.add_run(code_content)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif placeholder is not None:
            key = placeholder
            if key in math_blocks:
                latex = math_blocks[key]
                is_block = True
            else:
                latex = math_inlines.get(key, '')
                is_block = False
            png_path = os.path.join(tmpdir, f"math_{abs(hash(key))}.png")
            try:
                _render_latex_png(latex, png_path)
                run = p.add_run()
                width = Inches(4.5) if is_block else Inches(1.5)
                run.add_picture(png_path, width=width)
            except Exception:
                disp = f"$${latex}$$" if is_block else f"${latex}$"
                p.add_run(disp)
        else:
            p.add_run(full)

        pos = m.end()

    # Texte restant
    remaining = text[pos:]
    if remaining:
        p.add_run(remaining)


def _convert_table_line(line: str, doc, math_blocks, math_inlines, tmpdir):
    """Ajoute les lignes d'un tableau markdown au document."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn

    # Ignorer les lignes de séparation (---|---)
    if re.match(r'^[\|\s\-:]+$', line):
        return

    cells = [c.strip() for c in line.split('|')]
    cells = [c for c in cells if c]  # enlever les vides aux extrémités

    if not cells:
        return

    # Tableau simple : on crée un tableau d'une ligne
    table = doc.add_table(rows=1, cols=len(cells))
    table.style = 'Light Grid Accent 1'

    for i, cell_text in enumerate(cells):
        cell = table.cell(0, i)
        # Nettoyer le markdown inline dans la cellule
        cell_text_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
        cell_text_clean = re.sub(r'\*(.+?)\*', r'\1', cell_text_clean)
        cell_text_clean = re.sub(r'`(.+?)`', r'\1', cell_text_clean)
        cell.text = cell_text_clean


def create_docx_with_images(md_path: str, docx_path: str):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = Path(md_path).read_text(encoding='utf-8')
    text, math_blocks, math_inlines = _extract_placeholders(text)

    doc = Document()

    # Style par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    tmpdir = tempfile.mkdtemp(prefix='md2docx_')

    try:
        lines = text.splitlines()
        i = 0
        in_code_block = False
        code_lines = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Bloc de code (```...```)
            if stripped.startswith('```'):
                if in_code_block:
                    # Fin du bloc de code
                    if code_lines:
                        code_text = '\n'.join(code_lines)
                        p = doc.add_paragraph()
                        run = p.add_run(code_text)
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        p.paragraph_format.left_indent = Inches(0.3)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Ligne vide
            if not stripped:
                # doc.add_paragraph('')  # saut optionnel
                i += 1
                continue

            # Séparateur horizontal
            if stripped == '---' or stripped == '***' or stripped == '___':
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                pPr = p._p.get_or_add_pPr()
                from docx.oxml.ns import qn
                pBdr = pPr.makeelement(qn('w:pBdr'), {})
                bottom = pBdr.makeelement(qn('w:bottom'), {
                    qn('w:val'): 'single',
                    qn('w:sz'): '6',
                    qn('w:space'): '1',
                    qn('w:color'): '999999',
                })
                pBdr.append(bottom)
                pPr.append(pBdr)
                i += 1
                continue

            # Titres
            if stripped.startswith('# '):
                doc.add_heading(_clean_math_placeholders(stripped[2:], math_blocks, math_inlines), level=0)
                i += 1
                continue
            if stripped.startswith('## '):
                doc.add_heading(_clean_math_placeholders(stripped[3:], math_blocks, math_inlines), level=1)
                i += 1
                continue
            if stripped.startswith('### '):
                doc.add_heading(_clean_math_placeholders(stripped[4:], math_blocks, math_inlines), level=2)
                i += 1
                continue

            # Citation
            if stripped.startswith('> '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(stripped[2:])
                run.italic = True
                run.font.color.rgb = RGBColor(80, 80, 80)
                i += 1
                continue

            # Tableau (ligne avec |)
            if '|' in stripped:
                table_lines = []
                while i < len(lines) and '|' in lines[i].strip():
                    table_lines.append(lines[i])
                    i += 1
                # Construire le tableau
                _build_table(doc, table_lines)
                continue

            # Liste numérotée (1. 2. etc.)
            if re.match(r'^\d+\.\s', stripped):
                p = doc.add_paragraph()
                _add_formatted_run(p, stripped, math_blocks, math_inlines, tmpdir)
                p.paragraph_format.left_indent = Inches(0.3)
                i += 1
                continue

            # Liste à puces (- ou *)
            if stripped.startswith('- ') or stripped.startswith('* '):
                p = doc.add_paragraph()
                _add_formatted_run(p, stripped[2:], math_blocks, math_inlines, tmpdir)
                p.paragraph_format.left_indent = Inches(0.3)
                i += 1
                continue

            # Paragraphe normal
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped, math_blocks, math_inlines, tmpdir)
            i += 1

        doc.save(docx_path)
        return True

    finally:
        try:
            shutil.rmtree(tmpdir, ignore_errors=True)
        except Exception:
            pass


def _clean_math_placeholders(text, math_blocks, math_inlines):
    """Remplace les placeholders par leur rendu texte pour les titres."""
    for key, latex in math_blocks.items():
        text = text.replace(key, latex)
    for key, latex in math_inlines.items():
        text = text.replace(key, latex)
    return text


def _add_formatted_run(paragraph, text, math_blocks, math_inlines, tmpdir):
    """Ajoute du texte formaté dans un paragraphe existant."""
    from docx.shared import Inches, Pt

    # Gérer **bold**, *italic*, `code`, et placeholders math
    token_pat = re.compile(
        r'(\*\*(.+?)\*\*|'
        r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|'
        r'`(.+?)`|'
        r'(\x00(?:BLOCK|INLINE)\d+\x00))'
    )

    pos = 0
    for m in token_pat.finditer(text):
        before = text[pos:m.start()]
        if before:
            paragraph.add_run(before)

        full = m.group(0)
        bold_content = m.group(2)
        italic_content = m.group(3)
        code_content = m.group(4)
        placeholder = m.group(5)

        if bold_content is not None:
            run = paragraph.add_run(bold_content)
            run.bold = True
        elif italic_content is not None:
            run = paragraph.add_run(italic_content)
            run.italic = True
        elif code_content is not None:
            run = paragraph.add_run(code_content)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif placeholder is not None:
            key = placeholder
            latex = math_blocks.get(key) or math_inlines.get(key, '')
            is_block = key in math_blocks
            png_path = os.path.join(tmpdir, f"math_{abs(hash(key))}.png")
            try:
                _render_latex_png(latex, png_path)
                run = paragraph.add_run()
                width = Inches(4.5) if is_block else Inches(1.5)
                run.add_picture(png_path, width=width)
            except Exception:
                disp = f"$${latex}$$" if is_block else f"${latex}$"
                paragraph.add_run(disp)
        else:
            paragraph.add_run(full)

        pos = m.end()

    remaining = text[pos:]
    if remaining:
        paragraph.add_run(remaining)


def _build_table(doc, table_lines):
    """Construit un tableau python-docx à partir de lignes markdown."""
    from docx.shared import Inches, Pt

    # Filtrer les lignes de séparation
    data_lines = [l for l in table_lines if not re.match(r'^[\|\s\-:]+$', l.strip())]
    if not data_lines:
        return

    rows_data = []
    for line in data_lines:
        cells = [c.strip() for c in line.split('|')]
        cells = [c for c in cells if c]
        cleaned = [re.sub(r'\*\*(.+?)\*\*', r'\1', c) for c in cells]
        cleaned = [re.sub(r'\*(.+?)\*', r'\1', c) for c in cleaned]
        cleaned = [re.sub(r'`(.+?)`', r'\1', c) for c in cleaned]
        rows_data.append(cleaned)

    ncols = max(len(r) for r in rows_data)

    table = doc.add_table(rows=len(rows_data), cols=ncols)
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = table.cell(i, j)
                cell.text = cell_text
                # Première ligne = en-tête en gras
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph('')  # espacement après tableau


# ---------------------------------------------------------------------------
# Fonction principale
# ---------------------------------------------------------------------------

def convert_markdown_to_docx(md_path: str, docx_path: str):
    md_path = os.path.abspath(md_path)
    docx_path = os.path.abspath(docx_path)

    if not os.path.exists(md_path):
        print(f"ERREUR : fichier source '{md_path}' introuvable.")
        return False

    print(f"Conversion : {md_path} -> {docx_path}")

    # 1) pandoc (recommandé)
    if check_pandoc_installed():
        try:
            convert_with_pandoc(md_path, docx_path)
            print(f"[OK] Conversion réussie avec pandoc -> {docx_path}")
            return True
        except subprocess.CalledProcessError as e:
            print(f"[AVERTISSEMENT] pandoc a échoué : {e}")

    # 2) pypandoc
    try:
        import pypandoc
        convert_with_pypandoc(md_path, docx_path)
        print(f"[OK] Conversion réussie avec pypandoc -> {docx_path}")
        return True
    except ImportError:
        print("[INFO] pypandoc non installé (pip install pypandoc)")
    except Exception as e:
        print(f"[AVERTISSEMENT] pypandoc a échoué : {e}")

    # 3) Fallback matplotlib + python-docx
    try:
        import matplotlib
        from docx import Document
        print("[INFO] Utilisation du fallback matplotlib + python-docx...")
        ok = create_docx_with_images(md_path, docx_path)
        if ok:
            print(f"[OK] Conversion réussie avec fallback -> {docx_path}")
        return ok
    except ImportError as e:
        print(f"[ERREUR] Modules manquants pour le fallback : {e}")
        print("Installez : pip install matplotlib python-docx")
        return False


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    import argparse

    parser = argparse.ArgumentParser(
        description='Convertir un fichier Markdown avec équations LaTeX en .docx'
    )
    parser.add_argument(
        'source', nargs='?',
        default='article_scientifique.md',
        help='Fichier markdown source (défaut : article_scientifique.md)'
    )
    parser.add_argument(
        'dest', nargs='?',
        default=None,
        help='Fichier .docx destination (défaut : source avec extension .docx)'
    )
    args = parser.parse_args()

    src = args.source
    dst = args.dest or Path(src).with_suffix('.docx').name

    success = convert_markdown_to_docx(src, dst)
    sys.exit(0 if success else 1)
